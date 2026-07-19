#!/usr/bin/env python3
"""Generate the CLSSGB Learner Guide as BOTH a Markdown mirror (LG-*.md at repo
root) and a DOCX (courseware/LG-*.docx) from one source, so they never diverge.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt
body, one section per DMAIC phase, one sub-section per lab (Objective · Goal ·
What you'll build · Step-by-step · Check your work), plus quick-reference
formulas, assessment preparation and a glossary. All content is driven by
course_data + the domain data files, keeping the LG 100% aligned with the slide
deck, Lesson Plan and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
from data_domain5 import DOMAIN5
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4+DOMAIN5
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")
# Lab visuals are generated into the COURSE repo (courseware/assets/lg-visuals),
# which is a different place from the skill-local ASSETS dir used for the logos.
REPO_ASSETS=os.path.join(REPO,"courseware","assets")

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B=[]
def h1(t): B.append(("h1",t))
def h2(t): B.append(("h2",t))
def h3(t): B.append(("h3",t))
def p(t):  B.append(("p",t))
def bullets(xs): B.append(("bullets",xs))
def steps(xs): B.append(("steps",xs))
def img(path,caption): B.append(("img",path,caption))
def code(t): B.append(("code",t))
def note(t): B.append(("note",t))
def rule(): B.append(("rule",))

# ---------------- content ----------------
h1("Introduction")
p(f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
  "It follows the DMAIC roadmap end to end — Define, Measure, Analyze, Improve, Control — and provides "
  "step-by-step instructions for every hands-on lab. Core labs are assessed; elective labs are provided "
  "as additional practice and are run when time allows.")
p("The course content is grounded in the body of knowledge published by The Council for Six Sigma "
  "Certification (CSSC) in 'Six Sigma: A Complete Step-by-Step Guide', so what you learn here matches "
  "the recognised Green Belt standard (CSSC Green Belt scope, Chapters 1-24).")
p("Every lab uses one continuous scenario — the Northwind Retail Distribution Centre, where online orders "
  "are shipping late and customers are complaining. By the end of the course your lab outputs form a "
  "complete improvement package: project charter, VOC/CTQ, process and value stream maps, a validated "
  "measurement system, the baseline capability, proven root causes, selected countermeasures and a "
  "control plan.")

h1("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

h1("Skills Framework Alignment")
p(f"This course is aligned to the WSQ Technical Skills and Competencies (TSC) {C.TSC_TITLE} ({C.TSC_CODE}).")
h3("Abilities")
bullets(C.TSC_ABILITIES)
h3("Knowledge")
bullets(C.TSC_KNOWLEDGE)

h1("Before You Start")
h3("What you need")
bullets([
 "A laptop with a spreadsheet application (Microsoft Excel, Google Sheets or LibreOffice Calc).",
 "A browser, for the interactive problem-solving tools used in the Analyze labs.",
 "The course slides and this Learner Guide, downloaded from https://lms-tms.tertiaryinfotech.com.",
 "A work process of your own to think about — the tools apply far better when the example is real.",
])
h3("The interactive problem-solving toolkit")
p("Four browser-based tools are used during the labs. No installation or licence is required.")
bullets([
 "5 Whys — build and share a 5 Whys chain: https://alfredang.github.io/5whys/",
 "Fishbone Diagram — build an Ishikawa cause-and-effect diagram: https://alfredang.github.io/fishbone/",
 "Pareto Chart (collaborative) — your team brainstorms and votes in one live session and the Pareto chart builds itself: https://alfredang.github.io/paretochart/",
 "NovaSPC — run charts, SPC charts and process capability from your own CSV: https://alfredang.github.io/novaspc/",
])
h3("Core and elective labs")
bullets([
 "Core labs are completed by everyone and map directly to the assessment.",
 "Elective labs extend the same scenario with additional Lean Six Sigma tools; complete them if time allows or as post-course practice.",
 "All labs build on the same Northwind Retail Distribution Centre scenario, so outputs carry forward from one lab to the next.",
])
h3("Conventions used in every lab")
bullets([
 "Each lab states its objective, the deliverable you produce, the steps, and a check to confirm you are done.",
 "Tables shown in the steps can be built in a spreadsheet or on the worksheet provided.",
 "Where a lab uses an online tool, the tool URL is shown with the step.",
 "Keep every lab output — they combine into your final improvement package and are your revision material.",
])

# ---------------- per-topic, per-lab ----------------
for t in C.TOPICS:
    label = t["phase"].title() if t["num"] else "Foundations"
    h1(f"{t['phase']} — {t['title']}  ({t['weighting']})")
    p(t["subtitle"])
    h3("Key concepts")
    bullets([f"{name} — {desc}" for name, desc in t["concepts"]])
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        kind = "Elective" if a.get("elective") else "Core"
        title = a["title"].replace("Elective — ","")
        h2(f"Lab {a['num']} — {title}  [{kind}]")
        p(f"Objective: {a['objective']}")
        p(f"Goal: {a['desc']}")
        h3("What you'll build")
        p(a["build"]+f"   (Tools and techniques: {a['services']}.)")
        _vis=os.path.join(REPO_ASSETS,"lg-visuals",f"lab-{a['num']:02d}-visual.png")
        if os.path.exists(_vis):
            img(_vis,f"Lab {a['num']} at a glance — the deliverable, the tools and the steps.")
        h3("Step-by-step")
        steps([(instr,cmd) for instr,cmd in a["steps"]])
        h3("Check your work")
        p(a["test"])
        note(f"The full worksheet for this lab is in labs/lab-{a['num']:02d}-*.md.")
        rule()

h1("Quick Reference — Formulas You Should Know")
h3("Process performance metrics")
bullets([
 "Yield = (Good units / Total units) x 100",
 "DPU (Defects Per Unit) = Defects / Units",
 "DPO (Defects Per Opportunity) = Defects / (Units x Opportunities per unit)",
 "DPMO (Defects Per Million Opportunities) = DPO x 1,000,000",
 "First Pass Yield (FPY) = units passing with no rework / total units started",
 "Rolled Throughput Yield (RTY) = FPY(step 1) x FPY(step 2) x ... x FPY(step n)",
 "Process Cycle Efficiency = Value-added time / Total lead time",
 "Takt time = Available working time / Customer demand",
])
h3("Sigma level reference")
bullets([
 "1 sigma — 690,000 DPMO (about 31% yield)",
 "2 sigma — 308,000 DPMO (about 69% yield)",
 "3 sigma — 66,800 DPMO (about 93.3% yield)",
 "4 sigma — 6,210 DPMO (about 99.38% yield)",
 "5 sigma — 233 DPMO (about 99.977% yield)",
 "6 sigma — 3.4 DPMO (about 99.99966% yield)",
])

h1("Quick Reference — The Eight Wastes (DOWNTIME)")
bullets([
 "D — Defects: output that fails the requirement and must be corrected or redone.",
 "O — Overproduction: producing more, or earlier, than the customer needs.",
 "W — Waiting: work or people idle, waiting for the next step, an approval or information.",
 "N — Non-utilised talent: skills and ideas of people not being used.",
 "T — Transport: unnecessary movement of materials, work items or information between places.",
 "I — Inventory: work in progress, backlogs and queues sitting between steps.",
 "M — Motion: unnecessary movement of people, or switching between systems and screens.",
 "E — Extra-processing: doing more work to the output than the customer requires or values.",
])

h1("Preparing for the Assessment")
bullets([
 C.ASSESSMENT["written"],
 C.ASSESSMENT["practical"],
 "Both papers are open book — you may use these slides, this Learner Guide and your lab outputs.",
 "Revise by re-reading your own lab outputs; they follow exactly the same scenario as the Case Study.",
 "Be ready to define Lean, Six Sigma and Lean Six Sigma, and explain how they differ.",
 "Be ready to name the eight wastes and give a service-industry example of each.",
 "Be ready to explain each DMAIC phase, what it delivers and which tools belong to it.",
 "Be ready to calculate yield, DPU, DPO and DPMO from raw data and read off the sigma level.",
 "Be ready to explain how the Fishbone diagram and 5 Whys are used together to find a root cause.",
 "Re-work the labs from memory — being able to produce the tools unaided is the best preparation.",
 C.ASSESSMENT["note"],
])

h1("Glossary")
gl=[
 ("Lean","A method to maximise customer value by systematically identifying and removing waste."),
 ("Six Sigma","A data-driven method to reduce variation and defects; the target is 3.4 defects per million opportunities."),
 ("Lean Six Sigma","The combined method — Lean improves speed and flow, Six Sigma improves consistency and accuracy."),
 ("DMAIC","Define, Measure, Analyze, Improve, Control — the Six Sigma improvement roadmap."),
 ("PDCA","Plan, Do, Check, Act — a lighter improvement cycle used for small, fast improvements."),
 ("VOC","Voice of the Customer — customer needs and expectations expressed in the customer's own words."),
 ("CTQ","Critical to Quality — a specific, measurable requirement translated from a VOC statement."),
 ("SIPOC","Suppliers, Inputs, Process, Outputs, Customers — the macro 'as-is' process map."),
 ("Muda","The Japanese term for waste — any activity that consumes resource but creates no customer value."),
 ("DOWNTIME","Mnemonic for the eight wastes: Defects, Overproduction, Waiting, Non-utilised talent, Transport, Inventory, Motion, Extra-processing."),
 ("Value-added (VA)","An activity the customer would be willing to pay for because it changes the product or service."),
 ("Non-value-added (NVA)","An activity that consumes resource but adds nothing the customer values — pure waste."),
 ("Defect","An output that fails to meet the CTQ requirement."),
 ("Common cause variation","Natural, always-present variation built into the process; addressed by changing the process."),
 ("Special cause variation","Unusual variation traceable to a specific assignable event; addressed by investigating that event."),
 ("Pareto principle","The 80/20 rule — roughly 80% of effects come from 20% of causes."),
 ("Run chart","A plot of a metric over time, used to spot trend, shift, cluster and oscillation patterns."),
 ("5 Whys","Repeatedly asking 'why' to drill from a symptom down to an actionable root cause."),
 ("Fishbone (Ishikawa) diagram","A cause-and-effect diagram organising candidate causes into categories such as the 5Ms."),
 ("5M categories","Manpower, Method, Machine, Material, Measurement — standard Fishbone categories."),
 ("Multi-voting","A team technique to narrow a long list of candidate causes to an agreed shortlist."),
 ("DPU / DPO / DPMO","Defects per unit / per opportunity / per million opportunities — standard defect-rate metrics."),
 ("Yield","The percentage of output produced without defects."),
 ("RTY","Rolled Throughput Yield — the probability a unit passes through every process step with no rework."),
 ("Sigma level","A common yardstick for process performance, derived from DPMO."),
 ("COPQ","Cost of Poor Quality — the internal failure, external failure, appraisal and prevention costs of defects."),
 ("5S","Sort, Set in order, Shine, Standardise, Sustain — a Lean method for organising the workplace."),
 ("Poka-Yoke","Mistake proofing — designing the process so an error is difficult or impossible to make."),
 ("Standard work","The documented current best-known method for performing a task."),
 ("Kaizen","Continuous improvement through many small changes, made by everyone."),
 ("Takt time","Available working time divided by customer demand — the pace the process must run to meet demand."),
 ("Control plan","The document naming the metric, target, monitoring method, frequency, owner and reaction plan."),
 ("Specification limits (LSL/USL)","The acceptable limits set by the customer; outside them is a defect."),
 ("Control limits (LCL/UCL)","Limits calculated from the process itself, typically the mean plus/minus three standard deviations."),
 ("SPC","Statistical Process Control — monitoring a process over time to predict and prevent problems."),
 ("A3","A one-page report telling the whole improvement story from background through to follow-up."),
 ("Gemba","'The real place' — going to where the work actually happens to observe the process directly."),
]
B.append(("dl",gl))


# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)

def render_md():
    out=[f"# {C.TITLE} — Learner Guide",""]
    out.append(f"**WSQ Course Code:** {C.COURSE_CODE}  |  **Conducted by:** {C.ORG} ({C.UEN.replace('UEN: ','UEN ')})  |  **Version {C.VERSION} · {C.VERSION_DATE}**")
    out.append("")
    # TOC (h1 + h2)
    out.append("## Contents"); out.append("")
    for kind,*rest in B:
        if kind=="h1": out.append(f"- [{rest[0]}](#{_anchor(rest[0])})")
        elif kind=="h2": out.append(f"  - [{rest[0]}](#{_anchor(rest[0])})")
    out.append("")
    for kind,*rest in B:
        if kind=="h1": out+=["",f"## {rest[0]}",""]
        elif kind=="h2": out+=["",f"### {rest[0]}",""]
        elif kind=="h3": out+=[f"**{rest[0]}**",""]
        elif kind=="p": out+=[rest[0],""]
        elif kind=="bullets": out+=[f"- {x}" for x in rest[0]]+[""]
        elif kind=="steps":
            for i,(instr,cmd) in enumerate(rest[0],1):
                out.append(f"{i}. {instr}")
                if cmd: out+=["",f"   ```bash",f"   {cmd}","   ```",""]
            out.append("")
        elif kind=="img":
            _p,_cap=rest[0],rest[1]
            _rel=os.path.relpath(_p,REPO)
            out+=[f"![{_cap}]({_rel})","",f"*{_cap}*",""]
        elif kind=="code": out+=["```bash",rest[0],"```",""]
        elif kind=="note": out+=[f"> **Note:** {rest[0]}",""]
        elif kind=="rule": out+=["---",""]
        elif kind=="dl":
            for term,defn in rest[0]: out.append(f"- **{term}** — {defn}")
            out.append("")
    return "\n".join(out)

MD_OUT=os.path.join(REPO,f"LG-{C.SHORT_TITLE}.md")
with open(MD_OUT,"w") as f: f.write(render_md())
print("Saved",MD_OUT)

# ---------------- render DOCX ----------------
BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
INKCODE=RGBColor(0x0B,0x30,0x60)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("1","1 June 2026","Initial release - CLSSGB Learner Guide.",C.TRAINER),
 ("2","19 July 2026","Major revision: rebuilt as a 4-day Green Belt progression from the Yellow Belt "
  "(CLSSYB). Content deepened to the CSSC Green Belt scope (Chapters 1-24) - adds "
  "sampling and sample size calculation, Measurement System Analysis and Gage R&R, hypothesis testing "
  "with p-values and Type I/II errors, correlation and regression, FMEA and RPN, Design of Experiments, "
  "SPC control chart selection with the eight out-of-control rules, and process capability Cp/Cpk. "
  "Expanded from 10 to 25 hands-on labs across the five DMAIC phases; the running scenario is now the "
  "Northwind Retail Distribution Centre order-fulfilment process, matching the WA (SAQ) and PP "
  "assessment papers.",C.TRAINER),
 ("3",C.VERSION_DATE,"Removed the Practice Exam slide and its asset - there is no Six Sigma practice exam on exams.tertiaryinfotech.com.",C.TRAINER),
])
prodoc.add_toc(doc)

def code_para(text):
    for line in text.split("\n"):
        para=doc.add_paragraph(); prodoc._shade_para(para) if hasattr(prodoc,"_shade_para") else None
        r=para.add_run(line); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=INKCODE

for kind,*rest in B:
    if kind=="h1": doc.add_heading(rest[0],level=1)
    elif kind=="h2": doc.add_heading(rest[0],level=2)
    elif kind=="h3":
        para=doc.add_paragraph(); r=para.add_run(rest[0]); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=BRAND
    elif kind=="p": doc.add_paragraph(rest[0])
    elif kind=="bullets":
        for x in rest[0]: doc.add_paragraph(x,style="List Bullet")
    elif kind=="steps":
        # Write the step number as literal text rather than using Word's
        # "List Number" style: that style shares ONE counter for the whole
        # document, so Lab 9 would continue from Lab 8 (…62, 63, 64) instead of
        # restarting at 1. An explicit number restarts correctly for every lab.
        for i,(instr,cmd) in enumerate(rest[0],1):
            para=doc.add_paragraph()
            para.paragraph_format.left_indent=Pt(18)
            para.paragraph_format.first_line_indent=Pt(-18)
            r=para.add_run(f"{i}.  "); r.bold=True; r.font.color.rgb=BRAND
            para.add_run(instr)
            if cmd: code_para(cmd)
    elif kind=="img":
        _p,_cap=rest[0],rest[1]
        try:
            doc.add_picture(_p,width=Inches(6.2))
            doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
            cp=doc.add_paragraph(); cp.alignment=WD_ALIGN_PARAGRAPH.CENTER
            r=cp.add_run(_cap); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=GREY
        except Exception:
            pass
    elif kind=="code": code_para(rest[0])
    elif kind=="note":
        para=doc.add_paragraph(); r=para.add_run("Note: "); r.bold=True; r.font.color.rgb=BRAND
        para.add_run(rest[0]).font.size=Pt(10)
    elif kind=="rule": doc.add_paragraph("")
    elif kind=="dl":
        for term,defn in rest[0]:
            para=doc.add_paragraph(style="List Bullet")
            r=para.add_run(term+" — "); r.bold=True; para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
