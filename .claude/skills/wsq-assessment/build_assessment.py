#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the WSQ assessment set for 'Certified Lean Six Sigma Green Belt (CLSSGB) Training'
(TGS-2025055775):
  - Written Assessment (SAQ)  - 2 open-ended KNOWLEDGE questions (K1, K2), aligned to the slides
  - Practical Performance (PP) - 3 PRACTICAL tasks (A1-A5), aligned to the in-class labs
This MIRRORS the v5 reference papers exactly: same instrument, same question/task count, same
criterion codes and mapping, same timings (WA 60 min, PP 90 min, both open book). Only the content
is rewritten from this course's slides and labs, deepened to Green Belt level.
Each instrument is produced as a Question Paper and a matching Answer Key (4 DOCX total),
all with the WSQ house cover page (same as the Lesson Plan / Learner Guide). Body: Arial 11.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# This script lives in the wsq-assessment skill (.claude/skills/wsq-assessment/) and runs in
# place — it detects the course repo root by walking up to the nearest dir that has a .git
# folder (or both courseware/ and assessment/). Override with env REPO=/path if needed.
def _find_repo():
    env = os.environ.get("REPO")
    if env and os.path.isdir(env):
        return os.path.abspath(env)
    d = os.path.dirname(os.path.abspath(__file__))
    while d != os.path.dirname(d):
        if os.path.isdir(os.path.join(d, ".git")) or \
           (os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "assessment"))):
            return d
        d = os.path.dirname(d)
    return os.getcwd()

REPO = _find_repo()
# prodoc.py (WSQ cover page + version control + page numbers, same as LP/LG) ships with the
# tertiary-lesson-plan skill. Look for it at the project level first, then the user level.
for _cand in (os.path.join(REPO, ".claude/skills/tertiary-lesson-plan"),
              os.path.expanduser("~/.claude/skills/tertiary-lesson-plan")):
    if os.path.exists(os.path.join(_cand, "prodoc.py")):
        sys.path.insert(0, _cand); break
import prodoc  # cover page + version control + page numbers (same as LP/LG)

# ─── EDIT PER COURSE ────────────────────────────────────────────────────────
TITLE       = "Certified Lean Six Sigma Green Belt (CLSSGB) Training"
COURSE_CODE = "TGS-2025055775"
# ────────────────────────────────────────────────────────────────────────────
# The cover page renders prodoc's module-level TGS constant. Override it so the
# assessment cover shows THIS course's ref (works with either prodoc version —
# the older project prodoc has no course_code kwarg).
prodoc.TGS = f"TGS Ref No: {COURSE_CODE}"
OUT   = os.path.join(REPO, "assessment")

# Logos: prefer the course's own courseware/assets, else fall back to the copies bundled
# in this skill (so the assessment builds even outside this project). Replace the course
# logo per course; the Tertiary Infotech logo is the same for every WSQ course.
def _logo(name):
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(REPO, "courseware/assets", name), os.path.join(here, "assets", name)):
        if os.path.exists(p):
            return p
    return None
ORG_LOGO    = _logo("tertiary-infotech-logo.png")
COURSE_LOGO = _logo("clssgb-course-logo.png")   # None if absent → Tertiary-only cover (as LP/LG)

Q_VER, A_VER = "v6", "v6"   # single standardised version across all four files
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
# Assessments carry the cover page only — no Document Version Control Record.

# ---------------------------------------------------------------- WRITTEN (KNOWLEDGE)
# MIRRORS THE v5 PAPER: exactly 2 open-ended SAQ questions, criteria K1 and K2, 60 minutes.
# (criterion, context, question, [model-answer points]) — each traces to the course slides.
WRITTEN = [
 ("K1",
  "Northwind Retail Distribution Centre fulfils online orders from a 12,000 sqm warehouse. Customers have "
  "started complaining about late deliveries and the company is receiving negative feedback on social media. "
  "The Distribution Manager has asked a Green Belt to lead a structured improvement project rather than "
  "jumping straight to a solution.",
  "Describe the DMAIC methodology. For EACH of the five phases, state its purpose and give one key deliverable "
  "or tool the Green Belt would produce in that phase for the Northwind order-fulfilment project.",
  ["DMAIC is the disciplined, data-driven improvement roadmap used in Six Sigma: Define, Measure, Analyze, "
   "Improve, Control. Each phase ends in a tollgate review before the project proceeds.",
   "DEFINE - purpose: agree the problem, the customer requirement and the scope of the project. "
   "Deliverables/tools: project charter with a problem statement (process, time period, measurable gap and "
   "business impact, with no cause and no solution stated), goal statement, in-scope/out-of-scope table, "
   "Voice of the Customer, CTQ tree, SIPOC and a stakeholder/RACI map.",
   "MEASURE - purpose: quantify how the process actually performs today and prove the data can be trusted. "
   "Deliverables/tools: detailed process map or swimlane, value stream map, data collection plan with "
   "operational definitions, sampling plan and calculated sample size, Measurement System Analysis "
   "(Gage R&R), and the baseline metrics - yield, RTY, DPU, DPO, DPMO, sigma level and Cp/Cpk.",
   "ANALYZE - purpose: find and PROVE the root causes of the problem using data rather than opinion. "
   "Deliverables/tools: run chart and stability assessment, Pareto chart with stratification, Fishbone "
   "(5M+E), 5 Whys, and statistical validation via hypothesis testing (p-value against alpha), correlation "
   "and regression.",
   "IMPROVE - purpose: generate, select, de-risk and pilot solutions that address the proven root causes. "
   "Deliverables/tools: brainstorming/brainwriting and benchmarking, weighted solution selection matrix, "
   "5S, poka-yoke, standard work, FMEA with RPN scoring, Design of Experiments, and a pilot plan with "
   "success criteria and a rollback plan.",
   "CONTROL - purpose: hold the gain so the process does not drift back after the team disbands. "
   "Deliverables/tools: SPC control chart (selected for the data type), recalculated Cp/Cpk, control plan "
   "with metric, target, frequency, owner and reaction plan, SOP and standard work, visual management, "
   "and a formal handover to the process owner.",
   "Award the mark where the candidate names all five phases in the correct order, states a sensible purpose "
   "for each, and gives at least one relevant deliverable or tool per phase applied to the scenario. "
   "Accept any reasonable equivalent tool taught in the course."]),
 ("K2",
  "A large retail warehouse has been experiencing delays in its order fulfilment process. Customers are "
  "complaining about late deliveries and the company is receiving negative feedback on social media. The "
  "warehouse manager, Sarah, has decided to investigate the issue and identify the root causes of the delays. "
  "She plans to use several problem-solving techniques to analyse, prioritise and resolve the problem.",
  "List and briefly describe TWO problem-solving techniques that Sarah can use to identify, analyse and "
  "resolve the root causes of delays in the order fulfilment process at the warehouse. For each technique, "
  "explain how it works and what it would tell her about the delays.",
  ["Any TWO of the techniques taught in the course are acceptable. The candidate must name the technique, "
   "explain how it is applied, and say what it reveals about the delays.",
   "FISHBONE (ISHIKAWA / CAUSE-AND-EFFECT) DIAGRAM: the problem ('orders shipped late') is written in the "
   "fish head and candidate causes are organised onto bones by category - typically 5M+E: Manpower, Method, "
   "Machine, Material, Measurement and Environment. The team brainstorms causes onto each bone and asks 'why "
   "does this happen?' to add sub-causes. It gives Sarah a structured, categorised view of every plausible "
   "cause of the delays and prevents the team from fixating on one favourite theory.",
   "5 WHYS: starting from the symptom, ask 'why?' repeatedly (about five times) until an actionable process "
   "cause is reached. Each answer becomes the subject of the next why. The chain is validated by reading it "
   "backwards with 'therefore'. Stop when the answer is a process or system rather than a person - 'the "
   "picker was careless' is a symptom, not a root cause. It drills Sarah from 'the order was late' down to "
   "an underlying cause she can actually fix.",
   "PARETO ANALYSIS: the delay reasons are counted, sorted in descending order and plotted as bars with a "
   "cumulative percentage line. Applying the 80/20 rule, the categories to the left of the 80% crossing are "
   "the 'vital few'. It tells Sarah which small number of causes account for most of the late orders, so "
   "effort goes where it will have most effect. Stratifying by shift, carrier or product family shows "
   "whether the problem is universal or concentrated.",
   "5 WHYS / FISHBONE / PARETO / RUN CHARTS / HYPOTHESIS TESTING / MULTI-VOTING are all acceptable answers. "
   "Award the mark where the candidate names two distinct techniques, describes the mechanics of each "
   "correctly, and relates each to identifying or prioritising the causes of the order delays.",
   "A strong answer also notes that a candidate cause only becomes a proven ROOT cause once data supports "
   "it - for example by testing it with a hypothesis test and comparing the p-value against alpha."]),
]

# ---------------------------------------------------------------- PRACTICAL (ACTIVITY-BASED)
# MIRRORS THE v5 PAPER: exactly 3 tasks - Task 1 (A1, A2), Task 2 (A3), Task 3 (A4, A5), 90 minutes.
SCENARIO = (
 "A large retail warehouse has been experiencing significant delays in its order fulfilment process. "
 "Customers have started to complain about late deliveries, and the company is receiving negative feedback "
 "on social media. As a result, the warehouse manager, Sarah, has decided to apply the DMAIC methodology of "
 "Six Sigma to define the problem, measure current performance, and analyse the root causes of the delays. "
 "Sarah's goal is to create a clear project scope, collect key performance data, and identify the most "
 "critical issues affecting the order fulfilment process.\n\n"
 "Baseline data for the last month: 4,200 orders were shipped and 357 of them were delivered late "
 "(8.5%). The business wants late shipments reduced to 3.0%."
)

BOX_CAP = ("Submit your completed template to the assessor and paste a screenshot of your answer in the box "
           "below:")

PRACTICAL = [
 ("Task 1", "A1, A2",
  "DEFINE PHASE - Use a Project Charter Template to define the scope of the project. Include the following "
  "key elements in your project charter: project objective; problem statement; business impact; project scope "
  "(in-scope and out-of-scope activities); key stakeholders and team roles; project timeline and deliverables. "
  "ANALYZE PHASE - Use an Ishikawa (Fishbone) Diagram Template to identify potential causes of order delays. "
  "Organise the causes into at least four main categories (such as People, Process, Machines, Materials or "
  "Environment). "
  "(Lab 5 - Project Charter, Problem Statement, Goal Statement and Scope; Lab 6 - SIPOC, Stakeholder Analysis "
  "and RACI; Lab 16 - Fishbone, 5 Whys, Multi-Voting and Cause Prioritisation.)",
  BOX_CAP,
  "PROJECT CHARTER (Lab 5) - the candidate should produce a charter containing all six required elements:\n\n"
  "Project objective: a clear statement of what the project will achieve - e.g. 'Reduce late order shipments "
  "in the order fulfilment process so that customer delivery promises are met consistently.'\n\n"
  "Problem statement: must contain the four components - the process, the time period, the measurable gap and "
  "the business impact - and must name NO cause and NO solution. e.g. 'Over the last month, 357 of 4,200 "
  "orders shipped from the distribution centre (8.5%) were delivered later than the promised date, generating "
  "customer complaints and negative social media feedback.'\n\n"
  "Business impact: quantified where possible - cost of expedited freight, credits issued, rework/re-picking "
  "hours, lost repeat custom and reputational damage (the Cost of Poor Quality).\n\n"
  "Project scope: an explicit in-scope / out-of-scope table. In scope: order release through to carrier "
  "handover. Out of scope: e.g. carrier transit performance, returns processing, supplier inbound - anything "
  "the team cannot control. The process start and stop points must be stated.\n\n"
  "Key stakeholders and team roles: named sponsor, process owner, Green Belt, Black Belt mentor and SMEs. "
  "Accept a RACI with exactly one Accountable per deliverable.\n\n"
  "Project timeline and deliverables: a milestone schedule across the five DMAIC phases with a tollgate "
  "review closing each phase.\n\n"
  "A goal statement containing metric, baseline, target and date (e.g. 'reduce late shipments from 8.5% to "
  "3.0% by 31 December') is expected in a complete charter.\n\n"
  "FISHBONE DIAGRAM (Lab 16) - the effect 'Orders shipped late' in the fish head, with causes organised onto "
  "at least FOUR category bones. Using 5M+E, credible causes include:\n\n"
  "Manpower/People: insufficient pickers on the night shift; new staff not trained on the WMS; high absence; "
  "no cover at peak.\n"
  "Method/Process: no prioritisation of orders by carrier cut-off; picking route not optimised; batch rather "
  "than wave picking; manual re-keying between systems.\n"
  "Machine/Equipment: WMS or scanner downtime; conveyor stoppages; printer failures at the pack bench; "
  "forklift availability.\n"
  "Material: stock not in the mapped location; inaccurate inventory records; packaging shortages; damaged "
  "stock requiring replacement picks.\n"
  "Measurement: no agreed definition of 'late'; cut-off time recorded inconsistently; no visibility of "
  "order age.\n"
  "Environment: warehouse congestion at peak; seasonal volume spikes; aisle layout causing long travel.\n\n"
  "Award the mark where the charter carries all six required elements with a compliant problem statement, and "
  "the Fishbone has at least four populated category bones with plausible, specific causes - not one-word "
  "labels. Sub-causes on the major bones indicate a stronger answer."),

 ("Task 2", "A3",
  "IMPROVE PHASE - Use a Solution Selection Matrix Template to select the most feasible and impactful "
  "solution(s) from the list of improvement ideas generated for the order fulfilment process. Include the "
  "following key criteria in your selection process: Feasibility (Is it practical to implement?); Cost (Is it "
  "cost-effective?); Impact (Does it significantly reduce order delays?); Time to implement (How quickly can "
  "it be implemented?). Score every candidate solution against every criterion and rank the results. "
  "(Lab 19 - Solution Generation, Benchmarking and Brainwriting; Lab 21 - Solution Selection Matrix and "
  "Cost-Benefit Analysis.)",
  BOX_CAP,
  "SOLUTION SELECTION MATRIX (Lab 21) - the candidate should produce a matrix with candidate solutions down "
  "the rows and the four required criteria across the columns.\n\n"
  "Candidate solutions traceable to the Fishbone causes in Task 1, for example:\n"
  "- Sequence the pick list by carrier cut-off time so at-risk orders are picked first\n"
  "- Re-slot the top 20% fastest-moving SKUs closer to the pack bench\n"
  "- Introduce wave picking instead of picking order by order\n"
  "- Add a visual order-age board at the pack bench\n"
  "- Cross-train and add cover on the night shift\n"
  "- Preventive maintenance schedule for scanners and the conveyor\n"
  "- Cycle counting to correct inventory location accuracy\n\n"
  "Method the candidate should demonstrate:\n"
  "1. Assign a WEIGHT to each criterion reflecting what matters to the sponsor (impact and cost usually carry "
  "the highest weight).\n"
  "2. Agree the scoring scale BEFORE scoring, and define what a 1 and a 5 mean for each criterion so scores "
  "are comparable.\n"
  "3. Score each solution against each criterion, with the team rather than alone - divergent scores usually "
  "expose a hidden assumption worth discussing.\n"
  "4. Calculate the weighted score for each solution: multiply each score by its criterion weight and sum "
  "across the row.\n"
  "5. Rank the solutions by total weighted score and identify the top candidates.\n"
  "6. Cross-check the ranking on an effort-impact grid to spot the high-impact, low-effort quick wins.\n"
  "7. Confirm the selected solution addresses a PROVEN root cause and does not simply move the problem "
  "downstream.\n\n"
  "A complete answer states the selected solution(s) with the rationale, and ideally a cost-benefit view "
  "(implementation cost, ongoing cost, expected annual benefit and payback period).\n\n"
  "Award the mark where all four required criteria appear, every candidate solution is scored against every "
  "criterion, the scoring produces a defensible ranking, and the candidate names the selected solution(s) "
  "with a justification. Weighted scoring is expected at Green Belt level."),

 ("Task 3", "A4, A5",
  "CONTROL PHASE - Use a Control Plan Template to outline how the improved order fulfilment process will be "
  "monitored and controlled. The Control Plan should include: the key process metrics to be tracked (e.g. "
  "order processing time, number of late orders, picking time, shipping errors); data collection methods (Who "
  "will collect it? How often? Where will it be recorded?); control limits (What are the acceptable "
  "performance thresholds?); corrective actions to be taken if performance deviates from the control limits; "
  "and the roles and responsibilities of the team members in ensuring control. "
  "(Lab 23 - Statistical Process Control, Chart Selection and Control Limits; Lab 24 - Control Plan, SOP, "
  "Visual Management and Response Plan; Lab 25 - Verify the Gain, A3 Storyboard, Handover and Project "
  "Closure.)",
  BOX_CAP,
  "CONTROL PLAN (Lab 24) - the candidate should produce a table with one row per control point and the "
  "following columns: process step, metric (CTQ), specification/target, measurement method, sample size, "
  "frequency, owner and reaction plan.\n\n"
  "KEY PROCESS METRICS - the four named in the scenario, each with a target:\n"
  "- Order processing time (order release to carrier handover), e.g. target under 24 hours\n"
  "- Number/percentage of late orders, target 3.0% or below (down from the 8.5% baseline)\n"
  "- Picking time per order or per line\n"
  "- Shipping errors (wrong item, wrong address, wrong quantity)\n\n"
  "DATA COLLECTION METHOD - for each metric state WHO collects it, HOW OFTEN and WHERE it is recorded. e.g. "
  "'Shift supervisor; daily at end of shift; recorded in the WMS dashboard and displayed on the visual "
  "management board.' Operational definitions must match those used to collect the baseline, otherwise the "
  "before/after comparison is invalid.\n\n"
  "CONTROL LIMITS - the acceptable performance thresholds that trigger a response. At Green Belt level the "
  "candidate should distinguish CONTROL limits (calculated from the process itself at +/- 3 standard "
  "deviations - the voice of the process) from SPECIFICATION limits (set by the customer - the voice of the "
  "customer), and should NOT plot specification limits on a control chart. Selecting the correct chart for "
  "the data is expected: late orders counted daily with varying order volume is proportion-defective data "
  "with a varying sample size, which calls for a p-chart.\n\n"
  "CORRECTIVE ACTIONS - a specific reaction plan per metric stating what happens when the limit is breached "
  "and who decides. e.g. 'If late orders exceed 3% for two consecutive days, the shift supervisor escalates "
  "to the Distribution Manager, the pick sequence is re-prioritised by carrier cut-off, and the cause is "
  "logged for review at the daily huddle.' Reference to the out-of-control rules (a point beyond a control "
  "limit, nine points on one side of the centre line, six points trending) indicates a strong answer.\n\n"
  "ROLES AND RESPONSIBILITIES - named individual accountable for each control point, not a department. "
  "Typically: process owner accountable overall; shift supervisors collect and review daily; the Green Belt "
  "supports for the first 30/60/90 days; the sponsor reviews at the monthly business review.\n\n"
  "SUSTAINING ELEMENTS expected in a complete answer: an SOP/standard work document for the improved method; "
  "a visual management board; a daily team huddle; a training plan and record; and an audit schedule to "
  "verify the control plan is actually being followed.\n\n"
  "Award the mark where all five required elements are present, every control point has a NAMED owner, a "
  "monitoring frequency and a SPECIFIC reaction plan, and the candidate distinguishes control limits from "
  "specification limits."),
]

# ---------------------------------------------------------------- doc helpers
def base_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    return doc

def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p

def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)

def answer_box(doc, lines=None, code=None, height_pt=90):
    """1x1 bordered box. `lines` → bullet-style model answer; `code` → monospace
    code/YAML/command block (indentation preserved); neither → empty answer space."""
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    if code:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in code.split("\n"):
            b = cell.add_paragraph(style=None)
            b.paragraph_format.space_after = Pt(0); b.paragraph_format.space_before = Pt(0)
            rr = b.add_run(ln if ln else " ")
            rr.font.name = "Consolas"; rr.font.size = Pt(9)
            rr._element.rPr.rFonts.set(qn('w:cs'), "Consolas")
            wt = rr._element.find(qn('w:t'))
            if wt is not None: wt.set(qn('xml:space'), 'preserve')
    elif lines:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in lines:
            b = cell.add_paragraph(style=None); b.paragraph_format.left_indent = Inches(0.15)
            rr = b.add_run("•  " + ln); rr.font.size = Pt(10.5)
    else:
        # empty answer space
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr(); trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), str(int(height_pt*20))); trh.set(qn('w:hRule'), 'atLeast'); trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

FILL_GAP = 6    # extra space below each fill-in line (paired with double line spacing for writing room)

def candidate_block(doc):
    heading(doc, "Trainee Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

# Assessment briefing (from the course slides — "Briefing for Assessment").
BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]

LMS_URL = "https://lms-tms.tertiaryinfotech.com/"

def add_hyperlink(p, url, text):
    """Add a real clickable Word hyperlink (blue, underlined) to paragraph p."""
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)  # 11pt
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); p._p.append(link)
    return link

def instructions(doc, minutes_text):
    heading(doc, "Instructions to Candidate")
    # None marks the upload instruction, which carries a clickable LMS hyperlink.
    items = [
        "This is an individual exercise.",
        "This is an open-book assessment.",
        f"A total of {minutes_text} is given to complete this assessment.",
        None,
    ] + BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        if s is None:
            p.add_run(f"{i}.  Complete your answers on the document provided and "
                      "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)

def grading(doc, what):
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

def finish(doc, path):
    prodoc.add_page_numbers(doc); prodoc.enable_update_fields(doc)
    doc.save(path); print("  saved:", os.path.basename(path))

# ---------------------------------------------------------------- builders
def build_wa(answers):
    doc = base_doc()
    kind = "Written Assessment (SAQ) — Answer Key" if answers else "Written Assessment (SAQ)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Written Assessment (SAQ)" if answers else "Written Assessment (SAQ)",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; questions begin on the next page.
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has answered all written questions and demonstrated the underpinning "
                     "knowledge required for the course learning outcomes.")
        page_break(doc)
    para(doc, "Short-Answer Questions (Knowledge)", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Answer all questions in your own words. Each question tests underpinning knowledge covered in the "
              "course slides.", size=10.5, italic=True, color=GREY, after=8)
    # Pagination is EXPLICIT — two questions to a page on the paper, one model answer to a
    # page in the key. Do not swap this for Word's keepNext/cantSplit: Word pushes an
    # oversized box to the next page, but Google Docs draws the border anyway and prints the
    # question text and the page footer straight THROUGH it. See SKILL.md → Pagination.
    per_page = 1 if answers else 2
    for i, (crit, ctx, q, pts) in enumerate(WRITTEN, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, ctx, size=11, after=3)
        para(doc, f"{q}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=pts if answers else None)
        if i % per_page == 0 and i < len(WRITTEN):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to WA (SAQ) - {TITLE} - {suffix}.docx" if answers
            else f"WA (SAQ) - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

def build_pp(answers):
    doc = base_doc()
    kind = "Practical Performance (PP) — Answer Key" if answers else "Practical Performance (PP)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Practical Performance Assessment" if answers else "Practical Performance Assessment",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; the problem begins on the next page.
        candidate_block(doc); instructions(doc, "90 minutes")
        grading(doc, "Candidate has successfully completed all PP tasks and can explain the overall "
                     "functions and features used to achieve them.")
        page_break(doc)
    para(doc, "Practical Problem", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Scenario", size=11.5, bold=True, after=2)
    para(doc, SCENARIO, size=11, after=8)
    # Practical tasks are long and their boxes are tall, so they get a page each — on the
    # paper AND in the key. Same rule as the WA: the page break is ours, not the renderer's.
    for i, (label, crit, prompt, cap, pts) in enumerate(PRACTICAL, 1):
        para(doc, f"{label} ({crit}):", size=11.5, bold=True, after=2, before=6)
        para(doc, prompt, size=11, after=3)
        para(doc, cap, size=10.5, italic=True, color=GREY, after=4)
        answer_box(doc, code=pts if answers else None, height_pt=150)
        if i < len(PRACTICAL):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to PP Assessment - {TITLE} - {suffix}.docx" if answers
            else f"PP Assessment - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

if __name__ == "__main__":
    print("Building WSQ assessment set…")
    build_wa(answers=False); build_wa(answers=True)
    build_pp(answers=False); build_pp(answers=True)
    print(f"Done. WA: {len(WRITTEN)} questions · PP: {len(PRACTICAL)} tasks.")
